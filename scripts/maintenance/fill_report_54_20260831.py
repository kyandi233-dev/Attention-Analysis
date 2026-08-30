"""在 WPS 报告副本上填充 5.4 毫米波结果、表 6 三线表、图 11/12，并修正 5.1 毫米波口径。

用法: python scripts/maintenance/fill_report_54_20260831.py
输入: D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx (原地更新)
图源: D:/Project/厚粲杯/08_算法/output/06_正式实验/毫米波验证图_0831/
依赖: python-docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DOCX_PATH = Path(r"D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx")
FIG_DIR = Path(r"D:/Project/厚粲杯/08_算法/output/06_正式实验/毫米波验证图_0831")

FIGURES = {
    "5.4.1 数据可用范围与信号质量": (
        "MMWAVE_BLAND_ALTMAN_HR_2026-08-31.png",
        "图 11 毫米波融合心率与 ECG 金标准的对照",
        "注：图为 3 名佩戴 ECG 金标准参与者在 60 个探针前 30 s 窗口上，毫米波融合心率与 ECG 心率的 Bland-Altman 对照。",
    ),
    "5.4.2 呼吸、心搏相关变化与身体运动": (
        "MMWAVE_FUSION_MECHANISM_2026-08-31.png",
        "图 12 心率融合估计机制",
        "注：呈现时域峰间隔与频域谱峰两路心率估计及融合输出的机制。",
    ),
}

BODY = {
    "5.4.1 数据可用范围与信号质量": [
        "116 个实验场次中，110 个场次（58 名参与者）形成可用的毫米波采集数据，共 2,200 个探针窗口进入分析；其余 6 个场次的采集文件无法加载。逐帧毫米波数据按采集过程记录的实际时间与任务事件对应，连续分析窗口完全位于单个 Block 内。",
        "TABLE6",
        "以 3 名同时佩戴 ECG 与呼吸带的参与者在 60 个探针前 30 s 窗口上的对照显示，融合心率估计的窗口级 MAE 为 9.02 bpm，中位绝对误差 5.15 bpm，锁半频窗口占 2%；呼吸频率的 MAE 为 2.44 次/分，中位绝对误差 0.87 次/分，锁半频窗口占 8%。融合心率相对 ECG 呈约 4–6 bpm 的系统性低估，该对照仅覆盖 3 名参与者，属小样本探索性证据。",
        None,  # 图 11
        "单被试 5 场次的金标准对照曾发现，呼吸谐波落入心搏频带时可形成“强而错”的锁定：质量门控（心搏带信噪比、相位稳定性、时频差）全部通过而心率仍锁半频。当前融合估计以时域峰间隔与频域谱峰两路估计的融合纠正该问题：在上述 60 窗对照中，频域一路单独估计的 MAE 为 12.7–16.1 bpm，融合后降至 9.0–11.4 bpm；其中一名此前锁半频最严重的参与者，本次 25 s 融合估计的中位绝对误差为 1.13 bpm，锁半频窗口为 0。逐搏层面的对照显示，±75 ms 容差下逐搏识别的灵敏度为 0.170、精确率为 0.211，尚不足以支持逐搏间期与心率变异性指标的解释，因此本报告不报告相关结果。",
    ],
    "5.4.2 呼吸、心搏相关变化与身体运动": [
        "B1 与 B2 区块的融合心率中位数分别为 76.49 与 75.58 bpm，呼吸频率中位数均为 20.0 次/分，运动代理中位数均为 0.027，两个任务区块的毫米波特征水平接近。三种心率估计在场次级中位数上接近（融合 76.24、频域 72.46、时域 77.55 bpm）。",
        None,  # 图 12
    ],
    "5.4.3 毫米波特征与行为及思维探针的关系": [
        "探针前 30 s 窗口的关联分析纳入 58 名参与者、110 个实验场次和 2,200 个探针窗口，以参与者为聚类单位，预测变量在建模前标准化，系数表示预测变量每增加 1 个标准差对应的变化。",
        "Q2 警觉程度的有序 GEE 显示，呼吸频率每增加 1 个标准差对应更高的警觉等级，B = 0.079，95% CI [0.009, 0.150]；时间戳覆盖比例的对应系数为 B = 0.183，95% CI [0.049, 0.317]。融合、频域与时域心率的系数（B = −0.100、−0.115、−0.132）及运动代理、相位稳定性、心率置信度的系数 95% CI 均包含 0。",
        "Q1 注意内容的多项 Logistic 回归显示，融合心率每增加 1 个标准差，“大脑空白”相对“完全专注”的系数为 B = 0.534，95% CI [0.054, 1.015]；频域与时域心率的对应系数分别为 B = 0.497，95% CI [0.031, 0.962] 与 B = 0.536，95% CI [0.104, 0.968]，三个估计方向一致。运动代理与“走神”相对“完全专注”的系数为 B = 0.237，95% CI [0.050, 0.425]。“在任务上没想目标”相对参照类别的全部比较及呼吸频率与各类别的比较的 95% CI 均包含 0。",
        "窗口行为结局的高斯 GEE 显示，运动代理与 No-Go 误按率正相关，B = 1.449，95% CI [0.557, 2.342]，与 d′负相关，B = −3.374，95% CI [−5.505, −1.243]；心率置信度与总 Go 遗漏率（B = −0.013，95% CI [−0.023, −0.003]）、真遗漏率（B = −0.010，95% CI [−0.020, −0.001]）和 No-Go 误按率（B = −0.081，95% CI [−0.137, −0.026]）负相关，与 d′正相关（B = 0.235，95% CI [0.078, 0.393]）；呼吸频率与 No-Go 误按率负相关，B = −0.006，95% CI [−0.010, −0.001]。正确 Go RT 中位数与全部毫米波特征的系数 95% CI 均包含 0。",
        "参与者间与参与者内分量显示，呼吸频率与 Q2 的关联主要来自参与者内波动（参与者内 B = 0.062，95% CI [0.005, 0.120]；参与者间 95% CI 包含 0）。运动代理与 No-Go 误按率的关联在参与者间与参与者内两个分量上均不为 0（B = 0.051 与 0.020），呼吸频率的参与者内分量与 No-Go 误按率负相关（B = −0.019，95% CI [−0.036, −0.003]）。",
        "排除运动代理高于全样本第 90 百分位的窗口后（剩余 1,980 个窗口），心率与“大脑空白”的关联保持稳定（融合心率 B = 0.513，95% CI [0.001, 1.025]），呼吸频率与 Q2（B = 0.079，95% CI [0.005, 0.153]）及 No-Go 误按率（B = −0.006，95% CI [−0.011, −0.001]）的关联保持稳定；运动代理与“走神”、No-Go 误按率及 d′的关联在排除后 95% CI 包含 0。",
    ],
}

TABLE6_HEADER = ["特征", "M", "SD", "中位数", "范围"]
TABLE6_ROWS = [
    ["融合心率（bpm）", "76.48", "8.85", "76.24", "56.66–102.77"],
    ["频域心率（bpm）", "73.46", "10.32", "72.46", "52.09–102.54"],
    ["时域心率（bpm）", "78.41", "7.59", "77.55", "61.39–104.59"],
    ["呼吸频率（次/分）", "19.43", "3.23", "20.00", "8.26–26.15"],
    ["运动代理", "0.0281", "0.0079", "0.0276", "0.0143–0.0666"],
    ["相位稳定性", "0.9538", "0.0072", "0.9540", "0.9338–0.9702"],
    ["心率置信度", "0.3509", "0.1551", "0.3435", "0.0076–0.7459"],
    ["时间戳覆盖比例", "0.6408", "0.0775", "0.6363", "0.4879–0.7965"],
]
TABLE6_NOTE = "注：场次级统计先取每个实验场次内全部探针窗口的中位数，再对 110 个场次计算。"


def _set_run_fonts(run, *, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_text_para(anchor, text: str) -> None:
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Pt(21)
    pf.line_spacing = 1.5
    _set_run_fonts(p.add_run(text))


def _add_figure_para(anchor, fname: str, caption: str, note: str) -> None:
    p_img = anchor.insert_paragraph_before()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(FIG_DIR / fname), width=Cm(13.5))
    p_cap = anchor.insert_paragraph_before()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_fonts(p_cap.add_run(caption), bold=True)
    p_note = anchor.insert_paragraph_before()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_fonts(p_note.add_run(note), size_pt=9)


def _header_rule(table) -> None:
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tcPr.append(borders)
        b = borders.find(qn("w:bottom"))
        if b is None:
            b = OxmlElement("w:bottom")
            borders.append(b)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")


def _add_table6(anchor) -> None:
    doc = anchor._parent
    tbl = doc.add_table(rows=len(TABLE6_ROWS) + 1, cols=5, width=Cm(13.5))
    try:
        tbl.style = anchor.part.document.styles["三线表"]
    except KeyError:
        pass
    for j, h in enumerate(TABLE6_HEADER):
        tbl.rows[0].cells[j].text = h
    for i, row in enumerate(TABLE6_ROWS, start=1):
        for j, val in enumerate(row):
            tbl.rows[i].cells[j].text = val
    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    _set_run_fonts(run, size_pt=9, bold=(ri == 0))
    _header_rule(tbl)
    anchor._element.addprevious(tbl._element)


def main() -> None:
    doc = Document(str(DOCX_PATH))

    # 0) 修正 5.1 毫米波口径：111->110、5 场->6 场
    fixed = 0
    for p in doc.paragraphs:
        if "分别为" in p.text:
            for run in p.runs:
                if run.text.strip() == "111":
                    run.text = run.text.replace("111", "110")
                    fixed += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            if "毫米波" in row.cells[0].text and "未形成可用采集数据" in row.cells[-1].text:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if run.text.strip() == "111":
                                run.text = run.text.replace("111", "110")
                                fixed += 1
                            if run.text.strip().startswith("5") and "个场次" in run.text:
                                run.text = run.text.replace("5 ", "6 ", 1)
                                fixed += 1
    print(f"5.1 毫米波口径修正: {fixed} 处")

    # 1) 定位标题
    headings: dict[str, object] = {}
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in BODY or t in {"5.5 RGB 可见行为结果"}:
            headings[t] = p
    required = set(BODY) | {"5.5 RGB 可见行为结果"}
    if not required.issubset(headings):
        raise RuntimeError(f"缺失标题: {required - set(headings)}")

    anchors = {
        "5.4.1 数据可用范围与信号质量": headings["5.4.2 呼吸、心搏相关变化与身体运动"],
        "5.4.2 呼吸、心搏相关变化与身体运动": headings["5.4.3 毫米波特征与行为及思维探针的关系"],
        "5.4.3 毫米波特征与行为及思维探针的关系": headings["5.5 RGB 可见行为结果"],
    }

    # 2) 插入内容
    for heading, anchor in anchors.items():
        for item in BODY[heading]:
            if item is None:
                _add_figure_para(anchor, *FIGURES[heading])
            elif item == "TABLE6":
                _add_table6(anchor)
                p_cap = anchor.insert_paragraph_before()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_run_fonts(p_cap.add_run("表 6 毫米波特征的场次级描述统计（110 个实验场次）"), bold=True)
                p_note = anchor.insert_paragraph_before()
                p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_run_fonts(p_note.add_run(TABLE6_NOTE), size_pt=9)
            else:
                _add_text_para(anchor, item)

    doc.save(str(DOCX_PATH))
    print(f"已保存: {DOCX_PATH}")


if __name__ == "__main__":
    main()
