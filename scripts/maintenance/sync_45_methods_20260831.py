"""把 4.5 草稿（含方法引用）同步进 docx，并追加 11 条新文献到参考文献节。

用法: python scripts/maintenance/sync_45_methods_20260831.py
输入: D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx (原地更新)
草稿: FocusWave-Formal-Analysis/正式报告/章节草稿/4.5-数据处理与统计分析.md
依赖: python-docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

DOCX_PATH = Path(r"D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx")
DRAFT = Path(r"D:/Project/厚粲杯/08_算法/FocusWave-Formal-Analysis/正式报告/章节草稿/4.5-数据处理与统计分析.md")

NEW_REFS = [
    "Bazarevsky, V., Grishchenko, I., Raveendran, K., Zhu, T., Zhang, F., & Grundmann, M. (2020). BlazePose: On-device real-time body pose tracking. arXiv preprint arXiv:2006.10204.",
    "Chaudhary, A. K., Kothari, R., Acharya, M., Dangi, S., Nair, N., Bailey, R., Kanan, C., Diaz, G., & Pelz, J. B. (2020). RITnet: Real-time semantic segmentation of the eye for gaze tracking. In 2019 IEEE/CVF International Conference on Computer Vision Workshop (ICCVW) (pp. 3698–3702). IEEE. https://doi.org/10.1109/ICCVW.2019.00568",
    "Dragomiretskiy, K., & Zosso, D. (2014). Variational mode decomposition. IEEE Transactions on Signal Processing, 62(3), 531–544. https://doi.org/10.1109/TSP.2013.2288675",
    "Efron, B., & Tibshirani, R. J. (1994). An introduction to the bootstrap. CRC Press.",
    "Liang, K.-Y., & Zeger, S. L. (1986). Longitudinal data analysis using generalized linear models. Biometrika, 73(1), 13–22. https://doi.org/10.1093/biomet/73.1.13",
    "Lugaresi, C., Tang, J., Nash, H., McClanahan, C., Uboweja, E., Hays, M., Zhang, F., Chang, C.-L., Yong, M. G., Lee, J., Chang, W.-T., Hua, W., Georg, M., & Grundmann, M. (2019). MediaPipe: A framework for building perception pipelines. arXiv preprint arXiv:1906.08172.",
    "Macmillan, N. A., & Creelman, C. D. (2005). Detection theory: A user's guide (2nd ed.). Psychology Press.",
    "Redmon, J., Divvala, S., Girshick, R., & Farhadi, A. (2016). You only look once: Unified, real-time object detection. In Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (pp. 779–788). IEEE. https://doi.org/10.1109/CVPR.2016.91",
    "Sen, P. K. (1968). Estimates of the regression coefficient based on Kendall's tau. Journal of the American Statistical Association, 63(324), 1379–1389. https://doi.org/10.1080/01621459.1968.10480934",
    "Soukupová, T., & Čech, J. (2016). Real-time eye blink detection using facial landmarks. In 21st Computer Vision Winter Workshop (pp. 1–8).",
    "Theil, H. (1950). A rank-invariant method of linear and polynomial regression analysis. Indagationes Mathematicae, 12, 85–91.",
]


def _set_run_fonts(run, *, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _parse_draft() -> list[tuple[str | None, list[str]]]:
    """解析草稿：返回 [(小节标题或None, [段落...]), ...]。跳过章节标题与参考文献清单。"""
    text = DRAFT.read_text(encoding="utf-8")
    blocks: list[tuple[str | None, list[str]]] = []
    cur_title: str | None = None
    cur_paras: list[str] = []
    in_refs = False
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("## 本节参考文献"):
            in_refs = True
            continue
        if in_refs:
            continue
        if line.startswith("### "):
            if cur_paras or cur_title is not None:
                blocks.append((cur_title, cur_paras))
            cur_title = line[4:].strip()
            cur_paras = []
        elif line.startswith("## "):
            continue  # 章节级标题（4.5 已存在于 docx）
        elif line and not line.startswith("#"):
            cur_paras.append(line)
    if cur_paras or cur_title is not None:
        blocks.append((cur_title, cur_paras))
    return blocks


def main() -> None:
    doc = Document(str(DOCX_PATH))
    blocks = _parse_draft()
    print(f"草稿解析: {len(blocks)} 块")

    # 1) 定位 4.5 标题与 5 结果标题，删除之间的全部段落
    paras = doc.paragraphs
    i45 = next(i for i, p in enumerate(paras) if p.text.strip() == "4.5 数据分析")
    i5 = next(i for i, p in enumerate(paras) if p.text.strip() == "5 结果")
    anchor = paras[i5]
    for p in list(paras[i45 + 1:i5]):
        p._element.getparent().remove(p._element)
    print(f"已删除 4.5 旧正文 {i5 - i45 - 1} 段")

    # 2) 插入草稿内容（倒序插入块，块内正序）
    h4_style = None
    for p in doc.paragraphs:
        if p.text.strip() == "4.5.2 近红外瞳孔测量" or p.text.strip() in {t for t, _ in blocks if t}:
            pass
    # 取 5 结果标题样式作 Heading 3 参照，4.5 标题样式作 Heading 4 参照
    h45_style = paras[i45].style
    h5_style = paras[i5].style
    for title, content in blocks:
        if title is None:
            # 4.5 导言段
            for para_text in reversed(content):
                p = anchor.insert_paragraph_before(para_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(21)
                p.paragraph_format.line_spacing = 1.5
                _set_run_fonts(p.runs[0])
        else:
            for para_text in reversed(content):
                p = anchor.insert_paragraph_before(para_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(21)
                p.paragraph_format.line_spacing = 1.5
                _set_run_fonts(p.runs[0])
            h = anchor.insert_paragraph_before(title)
            h.style = h45_style
    print("4.5 草稿已同步")

    # 3) 参考文献节追加 11 条新文献（去重）
    refs_existing = {p.text.strip()[:60] for p in doc.paragraphs if p.text.strip()}
    added = 0
    ref_anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "参考文献":
            ref_anchor = p
    if ref_anchor is None:
        raise RuntimeError("未找到参考文献标题")
    tail = ref_anchor
    # 找参考文献节最后一段（在"8 附录"或其他标题之前）
    after_ref = None
    started = False
    for p in doc.paragraphs:
        if p.text.strip() == "参考文献":
            started = True
            continue
        if started and p.style.name.startswith(("Heading", "标题")) and p.text.strip():
            after_ref = p
            break
    anchor_ref = after_ref if after_ref is not None else ref_anchor
    for ref in reversed(NEW_REFS):
        if ref[:60] in refs_existing:
            continue
        p = anchor_ref.insert_paragraph_before(ref)
        p.paragraph_format.line_spacing = 1.5
        _set_run_fonts(p.runs[0])
        added += 1
    print(f"参考文献追加: {added} 条")

    doc.save(str(DOCX_PATH))
    print(f"已保存: {DOCX_PATH}")


if __name__ == "__main__":
    main()
