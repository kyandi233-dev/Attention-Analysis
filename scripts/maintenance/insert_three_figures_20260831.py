"""插入技术路线图、毫米波管线图、正式时间轴图，全文重编号。

用法: python scripts/maintenance/insert_three_figures_20260831.py
输入: D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx (原地更新)
图源: FocusWave-Formal-Analysis/正式报告/科研绘图/ 下三个 png
依赖: python-docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DOCX_PATH = Path(r"D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx")
FIG_DIR = Path(r"D:/Project/厚粲杯/08_算法/FocusWave-Formal-Analysis/正式报告/科研绘图")

NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"


def _set_run_fonts(run, *, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _has_image(p) -> bool:
    return bool(p._element.findall(f".//{NS}"))


def _insert_figure(anchor_after, fname: str, width_cm: float) -> None:
    p_img = anchor_after.insert_paragraph_before()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(FIG_DIR / fname), width=Cm(width_cm))


def _insert_caption(anchor_after, caption: str, note: str | None) -> None:
    p_cap = anchor_after.insert_paragraph_before()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_fonts(p_cap.add_run(caption), bold=True)
    if note:
        p_note = anchor_after.insert_paragraph_before()
        p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_fonts(p_note.add_run(note), size_pt=9)


def main() -> None:
    doc = Document(str(DOCX_PATH))

    # 1) 替换图 6（F02）为正式时间轴.png
    replaced = 0
    for i, p in enumerate(doc.paragraphs):
        if _has_image(p) and i + 1 < len(doc.paragraphs):
            if doc.paragraphs[i + 1].text.strip().startswith("图 6 正式实验范式"):
                for run in list(p.runs):
                    run._element.getparent().remove(run._element)
                p.add_run().add_picture(str(FIG_DIR / "正式时间轴.png"), width=Cm(14.0))
                replaced += 1
                break
    print(f"时间轴替换: {replaced}")

    # 2) 技术路线图插到 3 研究目标与研究问题标题后
    paras = doc.paragraphs
    idx3 = next(i for i, p in enumerate(paras) if p.text.strip() == "3 研究目标与研究问题")
    after3 = paras[idx3 + 1]
    _insert_figure(after3, "技术路线图_v1.png", 14.5)
    _insert_caption(after3, "图 1 FocusWave 多模态持续注意测评技术路线", None)
    print("技术路线图已插入")

    # 3) 毫米波管线图插到 4.5.4 标题后
    paras = doc.paragraphs
    idx454 = next(i for i, p in enumerate(paras) if p.text.strip() == "4.5.4 毫米波心肺相关活动与身体微动")
    after454 = paras[idx454 + 1]
    _insert_figure(after454, "毫米波管线图_v1.png", 14.5)
    _insert_caption(after454, "图 9 毫米波心率与呼吸估计处理管线",
                    "注：距离域复数信号在上游完成 Range FFT；时频融合用于纠正呼吸谐波引起的半频锁定；逐搏间期与心率变异性指标不进入正式分析。")
    print("毫米波管线图已插入")

    # 4) 全文重编号（实时列表；题注合并重写；引用一次性映射）
    paras = doc.paragraphs
    fig_idx = [i for i, p in enumerate(paras) if _has_image(p)]
    old2new: dict[int, int] = {}
    for new_no, idx in enumerate(fig_idx, start=1):
        if idx + 1 < len(paras):
            cap_p = paras[idx + 1]
            t = cap_p.text.strip()
            m = re.match(r"图 (\d+)(.*)", t)
            if m and len(t) < 60:
                old2new[int(m.group(1))] = new_no
                rest = m.group(2)
                for run in cap_p.runs:
                    run.text = ""
                cap_p.runs[0].text = f"图 {new_no}{rest}"
    ref_fixed = 0
    for p in paras:
        for run in p.runs:
            def repl(m):
                nonlocal ref_fixed
                n = int(m.group(1) or m.group(2))
                final = old2new.get(n, n)
                if final != n:
                    ref_fixed += 1
                return m.group(0).replace(str(n), str(final), 1)
            new_text = re.sub(r"见图 (\d+)|（图 (\d+)）", repl, run.text)
            if new_text != run.text:
                run.text = new_text
    print(f"重编号 {len(fig_idx)} 张；引用更新 {ref_fixed} 处")

    doc.save(str(DOCX_PATH))
    print(f"已保存: {DOCX_PATH}")


if __name__ == "__main__":
    main()
