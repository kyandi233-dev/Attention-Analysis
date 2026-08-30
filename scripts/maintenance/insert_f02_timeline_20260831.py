"""插入 F02 正式实验范式与时间轴图到 4.4，并全文顺延重编号。

用法: python scripts/maintenance/insert_f02_timeline_20260831.py
输入: D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx (原地更新)
图源: FocusWave-Formal-Analysis/正式报告/科研绘图/F02-正式实验范式与时间轴/rendered/f02-timeline.png
依赖: python-docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DOCX_PATH = Path(r"D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx")
F02 = Path(r"D:/Project/厚粲杯/08_算法/FocusWave-Formal-Analysis/正式报告/科研绘图/F02-正式实验范式与时间轴/rendered/f02-timeline.png")

NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
CAPTION_NOTE = "注：正式实验包含两个结构相同的 B 条件 Block；每个 Block 含 432 个试次（384 Go / 48 No-Go）与 10 次思维探针，探针后紧跟两问（Q1 注意内容、Q2 警觉程度）。单试次为刺激 250 ms + 掩蔽 900 ms。"


def _set_run_fonts(run, *, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _has_image(p) -> bool:
    return bool(p._element.findall(f".//{NS}"))


def main() -> None:
    doc = Document(str(DOCX_PATH))

    # 1) 定位 4.4 标题，其后插入 F02 图段+题注
    h44 = None
    for p in doc.paragraphs:
        if p.text.strip() == "4.4 正式实验程序与多模态采集":
            h44 = p
            break
    if h44 is None:
        raise RuntimeError("未找到 4.4 标题")
    anchor = h44
    # 找 4.4 标题后的第一个段落作插入锚点
    all_paras = doc.paragraphs
    idx44 = next(i for i, p in enumerate(all_paras) if p.text.strip() == "4.4 正式实验程序与多模态采集")
    after = all_paras[idx44 + 1]
    p_img = after.insert_paragraph_before()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(F02), width=Cm(14.0))
    p_cap = after.insert_paragraph_before()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_fonts(p_cap.add_run("图 6 正式实验范式与时间轴"), bold=True)
    p_note = after.insert_paragraph_before()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_fonts(p_note.add_run(CAPTION_NOTE), size_pt=9)
    print("F02 已插入")

    # 2) 全文重编号（实时列表；题注多 run 合并重写；引用一次性映射）
    paras = doc.paragraphs
    fig_idx = [i for i, p in enumerate(paras) if _has_image(p)]
    old2new: dict[int, int] = {}
    for new_no, idx in enumerate(fig_idx, start=1):
        if idx + 1 < len(paras):
            cap_p = paras[idx + 1]
            t = cap_p.text.strip()
            m = re.match(r"图 (\d+)(.*)", t)
            if m and len(t) < 60:
                old2new[int(m.group(1))] = new_no
                rest = m.group(2)
                for run in cap_p.runs:
                    run.text = ""
                cap_p.runs[0].text = f"图 {new_no}{rest}"
    print(f"重编号 {len(fig_idx)} 张，映射 {sorted(old2new.items())}")

    # 引用替换：一次性按最终映射（避免链式）
    ref_fixed = 0
    for p in paras:
        for run in p.runs:
            def repl(m):
                nonlocal ref_fixed
                n = int(m.group(1) or m.group(2))
                final = old2new.get(n, n)
                if final != n:
                    ref_fixed += 1
                return m.group(0).replace(str(n), str(final), 1)
            new_text = re.sub(r"见图 (\d+)|（图 (\d+)）", repl, run.text)
            if new_text != run.text:
                run.text = new_text
    print(f"引用更新: {ref_fixed} 处")

    doc.save(str(DOCX_PATH))
    print(f"已保存: {DOCX_PATH}")


if __name__ == "__main__":
    main()
