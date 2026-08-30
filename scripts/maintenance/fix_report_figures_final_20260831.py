"""报告副本图最终处理：打码替换、删附录图、补题注、全文重编号、插入 5.6 增量图。

用法: python scripts/maintenance/fix_report_figures_final_20260831.py
输入: D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx (原地更新)
依赖: python-docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DOCX_PATH = Path(r"D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx")
MASKED = Path(r"D:/Project/厚粲杯/08_算法/docs/交付/_docx_images/para_110_masked.jpeg")
FOREST = Path(r"D:/Project/厚粲杯/11_数据/_FormalAnalysis/MultiModal/full-20260831/figures/incremental_forest_56.png")

NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"


def _set_run_fonts(run, *, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_caption(anchor_para, caption: str, note: str) -> None:
    p_cap = anchor_para.insert_paragraph_before()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_fonts(p_cap.add_run(caption), bold=True)
    p_note = anchor_para.insert_paragraph_before()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_fonts(p_note.add_run(note), size_pt=9)


def _has_image(p) -> bool:
    return bool(p._element.findall(f".//{NS}"))


def main() -> None:
    doc = Document(str(DOCX_PATH))
    paras = doc.paragraphs

    # 1) 替换 4.3.1 被试照片为打码版（定位：图段后无题注且下一段为另一图段的第二张）
    # 直接按内容定位：4.3.1 中两个相邻图段（109/110 原索引），以文本上下文定位
    replaced = 0
    for i, p in enumerate(paras):
        if _has_image(p) and i + 1 < len(paras) and _has_image(paras[i + 1]):
            # 两个相邻图段：第一个是界面照片、第二个是被试照片
            for run in list(paras[i + 1].runs):
                run._element.getparent().remove(run._element)
            paras[i + 1].add_run().add_picture(str(MASKED), width=Cm(9.0))
            paras[i + 1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            replaced += 1
            break
    print(f"打码替换: {replaced}")

    # 2) 删除预实验轨迹图（4.3.1 6-block 与 4.3.2 3-block）与 4.5.1 RT 分布图
    removed = 0
    for p in list(paras):
        t = p.text.strip()
        if not _has_image(p):
            continue
        # 预实验轨迹图特征：图段前后是正文/标题且无题注；RT 分布图在 4.5.1（"平均反应时反映总体反应速度"段前）
        prev_t = paras[paras.index(p) - 1].text.strip() if paras.index(p) > 0 else ""
        next_t = paras[paras.index(p) + 1].text.strip() if paras.index(p) + 1 < len(paras) else ""
        if next_t.startswith("4.3.2") or next_t.startswith("4.4 "):
            p._element.getparent().remove(p._element)
            removed += 1
        elif next_t.startswith("平均反应时反映总体反应速度"):
            p._element.getparent().remove(p._element)
            removed += 1
    print(f"删除附录/QC 图: {removed}")

    # 3) 补题注（重编号前先给无题图补题注，编号在下一步统一）
    caps = {
        "实验任务界面示例": "注：任务界面以“榨汁大师”游戏情境呈现 SART 刺激。",
        "实验场景示例": "注：被试与多模态采集设备布置示意；图中人脸已作匿名化处理。",
        "RT-CV 的任务时间进程": "注：A 为 B1/B2 的场次配对均值；B 为 Block×cycle 的 RT-CV 轨迹（误差线为 95% CI）。",
    }
    added = 0
    for i, p in enumerate(paras):
        if not _has_image(p):
            continue
        nxt = paras[i + 1].text.strip() if i + 1 < len(paras) else ""
        if nxt.startswith("图 "):
            continue  # 已有题注
        if i + 1 < len(paras) and _has_image(paras[i + 1]):
            # 相邻图对：第一张（界面照片）题注插在第二张前；第二张（场景照）题注插在其后
            _add_caption(paras[i + 1], "图 4 实验任务界面示例", caps["实验任务界面示例"])
            after_second = paras[i + 2] if i + 2 < len(paras) else paras[i + 1]
            _add_caption(after_second, "图 5 实验场景示例", caps["实验场景示例"])
            added += 2
        elif nxt.startswith("5.2.2"):
            _add_caption(paras[i + 1], "图 9 RT-CV 的任务时间进程", caps["RT-CV 的任务时间进程"])
            added += 1
    print(f"补题注: {added}")

    # 4) 全文图重编号：按顺序收集图段，其后的"图 N"题注段按新序号更新
    paras = doc.paragraphs  # 增删后重新取实时列表
    fig_idx = [i for i, p in enumerate(paras) if _has_image(p)]
    print(f"当前图段数: {len(fig_idx)}")
    old2new: dict[int, int] = {}
    for new_no, idx in enumerate(fig_idx, start=1):
        if idx + 1 < len(paras):
            cap_p = paras[idx + 1]
            t = cap_p.text.strip()
            if t.startswith("图 ") and len(t) < 40:
                import re
                m = re.match(r"图 (\d+)", t)
                if m:
                    old2new[int(m.group(1))] = new_no
                    for run in cap_p.runs:
                        run.text = re.sub(r"^图 \d+", f"图 {new_no}", run.text)
    # 5) 正文引用同步
    ref_map = {}
    for old, new in old2new.items():
        ref_map[f"见图 {old}"] = f"见图 {new}"
        ref_map[f"（图 {old}）"] = f"（图 {new}）"
    fixed_refs = 0
    for p in paras:
        for run in p.runs:
            for old_s, new_s in ref_map.items():
                if old_s in run.text:
                    run.text = run.text.replace(old_s, new_s)
                    fixed_refs += 1
    print(f"重编号: {len(old2new)} 张；引用同步: {fixed_refs} 处")

    # 6) 插入 5.6 增量图（表 7 注之后、5.6.2 末尾文字之前）
    # 定位 5.6.3 标题作锚点，插在其前
    anchor56 = None
    for p in paras:
        if p.text.strip() == "5.6.3 模态互补、平均边际贡献与补充 AI 融合":
            anchor56 = p
            break
    if anchor56 is not None:
        p_img = anchor56.insert_paragraph_before()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(str(FOREST), width=Cm(14.5))
        p_cap = anchor56.insert_paragraph_before()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_fonts(p_cap.add_run(f"图 {len(fig_idx) + 1} 各模态组合相对行为基准的 ROC-AUC 增量"), bold=True)
        p_note = anchor56.insert_paragraph_before()
        p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_fonts(p_note.add_run("注：规则化多项 Logistic 回归，留一参与者交叉验证；误差棒为参与者级重抽样的 95% CI，红色为 95% CI 不含 0 的增量。"), size_pt=9)
        # 5.6.2 首段引用补充
        for p in paras:
            for run in p.runs:
                if "（表 7）" in run.text and "增量见图" not in run.text:
                    run.text = run.text.replace("（表 7）", f"（表 7，增量见图 {len(fig_idx) + 1}）")
        print("5.6 增量图已插入")

    doc.save(str(DOCX_PATH))
    print(f"已保存: {DOCX_PATH}")


if __name__ == "__main__":
    main()
