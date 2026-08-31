"""一次性完成报告副本全部规范修复（确定性脚本，逐步验证）。

用法: python scripts/maintenance/final_fix_all_20260831.py
覆盖: 摘要重写、引言平衡、2.4 同步、表格、参考文献重建、术语/口径/格式批量修复、图15/18替换
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

FW = r"D:/Project/厚粲杯/08_算法/FocusWave-Formal-Analysis/正式报告/0827报告v1_填充版_20260831.docx"
OUT = r"D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx"
DRAFT24 = Path(r"D:/Project/厚粲杯/08_算法/FocusWave-Formal-Analysis/正式报告/章节草稿/2.4-现有测评技术与无感连续传感.md")
NIR_DIR = Path(r"D:/Project/厚粲杯/11_数据/_FormalAnalysis/NIR/11_analysis_tables/figures/publication")
BLIP = ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"

NEW_ABSTRACT = ("持续注意并非由单次行为反应或单一生理信号直接可观测，其变化同时表现为行为、主观、眼部、心肺与外显行为的波动。"
                "研究以游戏化持续性注意反应任务为测量情境，经两阶段预实验确定任务负荷与刺激序列；"
                "61 名参与者完成 116 个实验场次，每场两个相同的任务区块，同步记录行为、思维探针、近红外瞳孔、毫米波与 RGB 视频。"
                "以行为与思维探针为认知参照，检验瞳孔几何、心率呼吸、可见行为特征与注意状态的效标关联，并在共同样本上比较各模态相对行为基准的预测增量。"
                "结果显示：瞳孔与 Go 遗漏及“大脑空白”报告相关，呼吸频率与警觉程度相关，眨眼候选事件率随任务进程升高并与较低警觉相关；"
                "各模态预测增量总体较小，其中 RGB 提供小幅稳定增益。"
                "行为与主观报告仍是当前最有效的测评参照，无感传感通道提供过程信息与连续监测可能。")

NEW_INTRO_22 = ("连续生理与外显行为测量为这一问题提供了补充路径，但不同技术具有不同边界。接触式生理设备能够获得较稳定的生理信号，却增加佩戴和操作负担；"
                "普通摄像头能够持续记录头部姿态和面部动作，但图像采集涉及隐私，并且外显行为与内部注意状态并不存在一一对应关系，其识别性能还依赖任务情境和可用特征（Bosch & D’Mello, 2021）；"
                "近红外眼部成像能够在主动照明下稳定记录瞳孔大小，但瞳孔同时受亮度、焦距、努力与唤醒调节，需要严格的视觉条件控制。")
NEW_INTRO_23 = ("毫米波（millimeter-wave, mmWave）雷达为低接触负担的连续测量提供了另一条值得检验的技术路径：调频连续波雷达利用人体回波中的相位和幅度变化记录胸壁及身体微动，"
                "由此提取呼吸、心搏及运动相关特征（Paterniani et al., 2023; Wang et al., 2021; Ge et al., 2024）。"
                "由于雷达采集的是电磁回波而非可识别面部图像，它在长时间、非接触和隐私敏感场景中具有潜在应用价值。")
NEW_INTRO_24 = ("然而，技术上的可测并不等同于心理测量上的有效，这一边界对上述无感传感通道同样成立。心率、呼吸和身体微动均会受到姿势、疲劳、情绪、运动以及个体差异等因素影响，它们并不是持续注意的特异性指标。"
                "现有毫米波研究更多回答“生命体征或动作能否被可靠提取”，而较少在标准化认知任务中使用独立的行为和主观效标检验“这些信号是否能够反映持续注意的变化”。"
                "同样，基于眨眼、点头等动作得到较高分类准确率，只能证明动作识别能力，不能直接说明已经获得了对持续注意构念的效度证据。"
                "将无感连续测量发展为注意状态测评手段，需要把信号提取、任务表现、主观体验和时间过程置于同一验证框架中。")

TABLE_ROWS = [
    ["测量方法", "主要可观察信息", "主要优势", "主要限制"],
    ["行为任务与思维探针", "反应表现、错误、反应时、主观注意内容与警觉", "与任务要求直接对应，解释路径清晰", "探针离散且依赖自我报告，行为结果可能具有多种形成机制"],
    ["EEG / fNIRS 等接触式生理", "神经电活动或血流动力学变化", "连续、客观，可研究神经机制", "佩戴负担、运动伪迹、设备和操作成本"],
    ["计算机视觉", "头姿、视线、面部动作、身体运动", "非接触、连续、设备易获得", "图像隐私、光照/遮挡/角度影响，外显行为不等同内部状态"],
    ["多模态融合", "多类行为、生理与主观信息", "可检验互补与增量信息", "同步、缺失、质量控制、模型复杂度和泛化问题"],
    ["毫米波雷达", "距离、速度、胸壁微动、身体微动及其派生生命体征", "非接触，不直接采集可识别面部图像，可连续感知微动", "运动、杂波、呼吸谐波和姿态影响明显；技术可测性不能替代注意测量效度验证"],
]

METHOD_REFS = [
    "Bazarevsky, V., Grishchenko, I., Raveendran, K., Zhu, T., Zhang, F., & Grundmann, M. (2020). BlazePose: On-device real-time body pose tracking. arXiv preprint arXiv:2006.10204.",
    "Chaudhary, A. K., Kothari, R., Acharya, M., Dangi, S., Nair, N., Bailey, R., Kanan, C., Diaz, G., & Pelz, J. B. (2020). RITnet: Real-time semantic segmentation of the eye for gaze tracking. In 2019 IEEE/CVF International Conference on Computer Vision Workshop (ICCVW) (pp. 3698–3702). IEEE. https://doi.org/10.1109/ICCVW.2019.00568",
    "Dragomiretskiy, K., & Zosso, D. (2014). Variational mode decomposition. IEEE Transactions on Signal Processing, 62(3), 531–544. https://doi.org/10.1109/TSP.2013.2288675",
    "Efron, B., & Tibshirani, R. J. (1994). An introduction to the bootstrap. CRC Press.",
    "Liang, K.-Y., & Zeger, S. L. (1986). Longitudinal data analysis using generalized linear models. Biometrika, 73(1), 13–22. https://doi.org/10.1093/biomet/73.1.13",
    "Lugaresi, C., Tang, J., Nash, H., McClanahan, C., Uboweja, E., Hays, M., Zhang, F., Chang, C.-L., Yong, M. G., Lee, J., Chang, W.-T., Hua, W., Georg, M., & Grundmann, M. (2019). MediaPipe: A framework for building perception pipelines. arXiv preprint arXiv:1906.08172.",
    "Macmillan, N. A., & Creelman, C. D. (2005). Detection theory: A user's guide (2nd ed.). Psychology Press.",
    "Redmon, J., Divvala, S., Girshick, R., & Farhadi, A. (2016). You only look once: Unified, real-time object detection. In Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (pp. 779–788). IEEE. https://doi.org/10.1109/CVPR.2016.91",
    "Sen, P. K. (1968). Estimates of the regression coefficient based on Kendall's tau. Journal of the American Statistical Association, 63(324), 1379–1389. https://doi.org/10.1080/01621459.1968.10480934",
    "Soukupová, T., & Čech, J. (2016). Real-time eye blink detection using facial landmarks. In 21st Computer Vision Winter Workshop (pp. 1–8).",
    "Theil, H. (1950). A rank-invariant method of linear and polynomial regression analysis. Indagationes Mathematicae, 12, 85–91.",
]
EXTRA_REFS = [
    "Baltrušaitis, T., Ahuja, C., & Morency, L.-P. (2019). Multimodal machine learning: A survey and taxonomy. IEEE Transactions on Pattern Analysis and Machine Intelligence, 41(2), 423–443. https://doi.org/10.1109/TPAMI.2018.2798607",
    "Campbell, D. T., & Fiske, D. W. (1959). Convergent and discriminant validation by the multitrait-multimethod matrix. Psychological Bulletin, 56(2), 81–105. https://doi.org/10.1037/h0046016",
    "Cronbach, L. J., & Meehl, P. E. (1955). Construct validity in psychological tests. Psychological Bulletin, 52(4), 281–302. https://doi.org/10.1037/h0040957",
    "Curran, P. J., & Bauer, D. J. (2011). The disaggregation of within-person and between-person effects in longitudinal models of change. Annual Review of Psychology, 62, 583–619. https://doi.org/10.1146/annurev.psych.093008.100356",
    "D'Mello, S. K., & Kory, J. (2015). A review and meta-analysis of multimodal affect detection systems. ACM Computing Surveys, 47(3), Article 43, 1–36. https://doi.org/10.1145/2682899",
    "Smilek, D., Carriere, J. S. A., & Cheyne, J. A. (2010). Out of mind, out of sight: Eye blinking as indicator and embodiment of mind wandering. Psychological Science, 21(6), 786–789. https://doi.org/10.1177/0956797610368063",
]


def set_fonts(run, size=10.5, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def has_img(p):
    return bool(p._element.findall(BLIP))


def n_imgs(doc):
    return sum(1 for p in doc.paragraphs if has_img(p))


def replace_run_text(p, old, new):
    """在非图 run 上做文本替换。"""
    n = 0
    for run in p.runs:
        if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
            continue
        if old in run.text:
            run.text = run.text.replace(old, new)
            n += 1
    return n


def main():
    for DOCX in (FW, OUT):
        doc = Document(DOCX)
        paras = doc.paragraphs
        assert n_imgs(doc) == 25, f"起始图段 != 25: {n_imgs(doc)}"

        print(f"  # 1) 摘要 图段={n_imgs(doc)}", flush=True)
# 1) 摘要 + 引言 3 段
        for p in paras:
            t = p.text.strip()
            if t.startswith("摘要 ") and "游戏化持续性注意" in t:
                for run in p.runs:
                    run.text = ""
                p.runs[0].text = "摘要 " + NEW_ABSTRACT
            elif t.startswith("连续生理与外显行为测量为这一问题提供了补充路径") and "近红外" not in t:
                for run in p.runs:
                    run.text = ""
                p.runs[0].text = NEW_INTRO_22
            elif t.startswith("毫米波（millimeter-wave") and "另一条" not in t:
                for run in p.runs:
                    run.text = ""
                p.runs[0].text = NEW_INTRO_23
            elif t.startswith("然而，技术上的可测并不等同于心理测量上的有效") and "无感传感通道" not in t:
                for run in p.runs:
                    run.text = ""
                p.runs[0].text = NEW_INTRO_24

        print(f"  # 2) 2.4 图段={n_imgs(doc)}", flush=True)
# 2) 2.4 整节替换（标题在块首、正文正序）
        paras = doc.paragraphs
        i24 = next(i for i, p in enumerate(paras) if p.text.strip().startswith("2.4 "))
        i25 = next(i for i, p in enumerate(paras) if p.text.strip().startswith("2.5 "))
        for p in list(paras[i24:i25]):
            p._element.getparent().remove(p._element)
        d24 = DRAFT24.read_text(encoding="utf-8")
        title24 = None
        body24 = []
        for line in d24.splitlines():
            s = line.strip()
            if s.startswith("## 本节参考文献"):
                break
            if s.startswith("# "):
                title24 = s.lstrip("#").strip()
            elif s and not s.startswith("|"):
                body24.append(s)
        paras = doc.paragraphs
        anchor = next(p for p in paras if p.text.strip().startswith("2.5 "))
        h_style = anchor.style
        for para_text in body24:  # 正序
            np_ = anchor.insert_paragraph_before()
            np_.alignment = 3
            np_.paragraph_format.first_line_indent = Pt(21)
            np_.paragraph_format.line_spacing = 1.5
            set_fonts(np_.add_run(para_text))
        h = anchor.insert_paragraph_before(title24)
        h.style = h_style
        # 2.4 比较表（插到"因此，现有测量方法更适合"段后）
        paras = doc.paragraphs
        anchor_tbl = next(p for p in paras if p.text.strip().startswith("因此，现有测量方法更适合"))
        tbl = doc.add_table(rows=len(TABLE_ROWS), cols=4)
        try:
            tbl.style = anchor_tbl.part.document.styles["三线表"]
        except KeyError:
            pass
        for ri, row in enumerate(TABLE_ROWS):
            for ci, val in enumerate(row):
                cell = tbl.rows[ri].cells[ci]
                cell.text = val
                for cp in cell.paragraphs:
                    cp.alignment = 1
                    for run in cp.runs:
                        set_fonts(run, size=9, bold=(ri == 0))
        for cell in tbl.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            b = OxmlElement("w:tcBorders")
            tcPr.append(b)
            bt = OxmlElement("w:bottom")
            bt.set(qn("w:val"), "single")
            bt.set(qn("w:sz"), "6")
            bt.set(qn("w:space"), "0")
            bt.set(qn("w:color"), "auto")
            b.append(bt)
        anchor_tbl._element.addprevious(tbl._element)
        p_cap = anchor_tbl.insert_paragraph_before()
        p_cap.alignment = 1
        set_fonts(p_cap.add_run("表 1 常见持续注意测量方法比较"), bold=True)

        print(f"  # 3) 表重编号 图段={n_imgs(doc)}", flush=True)
# 3) 表重编号 + 引用
        paras = doc.paragraphs
        tab_paras = [(i, p) for i, p in enumerate(paras) if re.match(r"表 \d+", p.text.strip()) and len(p.text.strip()) < 55]
        old2new = {}
        for new_no, (i, p) in enumerate(tab_paras, start=1):
            m = re.match(r"表 (\d+)(.*)", p.text.strip())
            old2new[int(m.group(1))] = new_no
            for run in p.runs:
                run.text = ""
            p.runs[0].text = f"表 {new_no}{m.group(2)}"
        for p in paras:
            for run in p.runs:
                if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
                    continue
                def repl(m, old2new=old2new):
                    n = int(m.group(1))
                    final = old2new.get(n, n)
                    return m.group(0).replace(str(n), str(final), 1)
                run.text = re.sub(r"见表 (\d+)", repl, run.text)
                run.text = re.sub(r"（表 (\d+)", repl, run.text)

        print(f"  # 4) 批量文本修复 图段={n_imgs(doc)}", flush=True)
# 4) 批量文本修复
        TEXT_FIXES = [
            ("正式实验由 63 名参与者完成 119 次实验", "正式实验由 61 名参与者完成 116 次实验"),
            ("和 10 次双问题探针", "和 20 次双问题探针"),
            ("控制被试的被打断感", "控制参与者的被打断感"),
            ("被试需要全程尽量保持直坐", "参与者需要全程尽量保持直坐"),
            ("被试为定向选择", "参与者为定向选择"),
            ("18个带反馈的练习试次", "18 个带反馈的练习试次"),
            ("Block 间2分钟休息", "Block 间 2 分钟休息"),
            ("正确Go试次反应时", "正确 Go 试次反应时"),
            ("图2和图3均为预实验阶段二和正式实验双问题探针流程示意图", "图 3 和图 4 均为预实验阶段二和正式实验双问题探针流程示意图"),
            ("且本研究未对瞳孔的参与者间与参与者内效应在该模型内分解", "探针层模型未单独报告瞳孔两个分量的系数"),
            ("本研究未对瞳孔的参与者间与参与者内效应在该模型内分解", "探针层模型未单独报告瞳孔两个分量的系数"),
            ("不宜简单排列为高低等级", "更适合按观察层面与限制条件区分"),
            ("再评估以多模态为教师蒸馏无感单模态的路线", "再评估以多模态信号辅助训练无感单模态模型的路线"),
            ("图为 3 名佩戴 ECG 金标准参与者在 60 个探针前 30 s 窗口上", "图为 5 名佩戴 ECG 金标准参与者中取 3 名共 60 个探针前 30 s 窗口"),
            ("因该参与者类别不全而无法定义，表中为该指标", "因该参与者测试折中类别不全而无法定义 ROC-AUC，表中为该指标"),
            ("呼吸谐波引起的半频锁定", "呼吸谐波引起的锁半频"),
            ("图 3 预实验流程 (A) 各试验类型的时间结构与关键时间节点", "图 3 试次时间结构与关键时间节点"),
            ("图 4 实验流程 (B) 屏幕呈现的视觉刺激序列", "图 4 视觉刺激序列"),
        ]
        for p in paras:
            for old, new in TEXT_FIXES:
                replace_run_text(p, old, new)
        # 表格内（表 5 术语、表 4 毫米波场次数）
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = (
                                run.text.replace("无时序歧义 Go 遗漏率", "真遗漏率")
                                .replace("时序歧义 Go 遗漏率", "预判遗漏率")
                                .replace("无时序歧义遗漏率", "真遗漏率")
                                .replace("时序歧义遗漏率", "预判遗漏率")
                                .replace("5 个场次未形成可用采集数据", "6 个场次未形成可用采集数据")
                            )

        print(f"  # 5) 重复指引段 图段={n_imgs(doc)}", flush=True)
# 5) 重复指引段去重
        for key in [
            "毫米波外部数据集验证、双机同步校准与文献算法重现的完整记录见附录 D。",
            "RGB 全部 96 行模型效应见附录 E。",
            "八组合的 61 折逐折性能明细见附录 E。",
        ]:
            hits = [p for p in doc.paragraphs if p.text.strip() == key]
            for p in hits[1:]:
                p._element.getparent().remove(p._element)

        print(f"  # 6) 平台图号 图段={n_imgs(doc)}", flush=True)
# 6) 平台图号 26/27（跳过含图 run）
        for p in doc.paragraphs:
            replace_run_text(p, "图 25 无感测评原型平台界面（网页端）", "图 26 无感测评原型平台界面（网页端）")
            replace_run_text(p, "图 26 无感测评原型平台界面（移动端）", "图 27 无感测评原型平台界面（移动端）")

        print(f"  # 7) 正文题目 图段={n_imgs(doc)}", flush=True)
# 7) 正文题目统一居中
        for p in doc.paragraphs:
            if p.text.strip().startswith("FocusWave：多模态持续注意测评工具"):
                for run in p.runs:
                    run.text = "FocusWave：面向注意状态的无感多模态智能测评系统"
                p.alignment = 1
                p.paragraph_format.first_line_indent = None

        print(f"  # 8) 特色创新 图段={n_imgs(doc)}", flush=True)
# 8) 特色创新应用段并入正文
        paras = doc.paragraphs
        i7 = next(i for i, p in enumerate(paras) if p.text.strip() == "7 特色与创新")
        for i in range(i7 + 1, len(paras)):
            t = paras[i].text.strip()
            if t.startswith("在应用层面"):
                app_text = ("应用层面已形成可运行的无感测评原型平台：毫米波数据采集、状态可视化、报告生成与移动端适配均已贯通，"
                            "并预留多模态特征与预测模型接口；该平台当前定位为数据采集、可视化与模型接入的原型，其测评解释力受限于前述效度证据的边界。")
                prev = paras[i - 1]
                old_text = "".join(r.text for r in prev.runs)
                for run in prev.runs:
                    run.text = ""
                prev.runs[0].text = old_text.rstrip("。") + "。" + app_text
                paras[i]._element.getparent().remove(paras[i]._element)
                break

        print(f"  # 9) 6.2 空段 图段={n_imgs(doc)}", flush=True)
# 9) 6.2 空段清理 + 过渡句
        paras = doc.paragraphs
        i62 = next(i for i, p in enumerate(paras) if p.text.strip() == "6.2 多模态增量、测量边界与未来方向")
        for p in list(paras[i62 + 1 : i62 + 8]):
            if not p.text.strip():
                p._element.getparent().remove(p._element)
        paras = doc.paragraphs
        i62 = next(i for i, p in enumerate(paras) if p.text.strip() == "6.2 多模态增量、测量边界与未来方向")
        if paras[i62 + 1].text.strip().startswith("未来方向包括"):
            new_p = paras[i62 + 1].insert_paragraph_before()
            new_p.alignment = 3
            new_p.paragraph_format.first_line_indent = Pt(21)
            set_fonts(new_p.add_run("多模态增量与测量边界已在上述结果中呈现，本节在此基础上说明局限与后续方向。"))

        print(f"  # 10) 参考文献 图段={n_imgs(doc)}", flush=True)
# 10) 参考文献重建（从另一份收集原条目 + METHOD + EXTRA，去重字母序）
        paras = doc.paragraphs
        entries = []
        for i, p in enumerate(paras):
            t = p.text.strip()
            if not t or p.style.name.startswith("Heading"):
                continue
            if t.startswith(("注：", "图 26", "图 27", "图 25")):
                continue
            if re.match(r"^[A-ZÀ-ÝČĎŽ]", t) and "(" in t and ")" in t and "（" not in t[:20]:
                entries.append(t)
                p._element.getparent().remove(p._element)
        all_refs = entries + METHOD_REFS + EXTRA_REFS
        seen = set()
        uniq = []
        for r in all_refs:
            key = r[:60]
            if key not in seen:
                seen.add(key)
                uniq.append(r)

        def sort_key(r):
            m = re.match(r"^[A-ZÀ-ÝČĎŽ'’\-]+", r)
            return (m.group(0).lower() if m else r.lower())

        uniq_sorted = sorted(uniq, key=sort_key)
        ref_title = next(p for p in doc.paragraphs if p.text.strip() == "参考文献")
        for r in reversed(uniq_sorted):
            p = ref_title.insert_paragraph_before()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = -266700
            run = p.add_run(r)
            run.font.size = Pt(10.5)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        print(f"  # 11) 图 15/18 图段={n_imgs(doc)}", flush=True)
# 11) 图 15/18 替换为 1×3 新版
        MAP_REPLACE = {
            "图 15 瞳孔数据质量与覆盖": NIR_DIR / "Figure09_data_quality_coverage.png",
            "图 18 瞳孔效应估计的稳健性": NIR_DIR / "Figure10_robustness_models.png",
        }
        from docx.shared import Cm
        paras = doc.paragraphs
        for i, p in enumerate(paras):
            if not has_img(p):
                continue
            nxt = paras[i + 1].text.strip() if i + 1 < len(paras) else ""
            if nxt in MAP_REPLACE:
                for run in list(p.runs):
                    run._element.getparent().remove(run._element)
                p.add_run().add_picture(str(MAP_REPLACE[nxt]), width=Cm(14.5))
                p.alignment = 1

        doc.save(DOCX)
        doc2 = Document(DOCX)
        n_after = n_imgs(doc2)
        refs_after = sum(1 for p in doc2.paragraphs if re.match(r"^[A-ZÀ-ÝČĎŽ]", p.text.strip()) and "(" in p.text and ")" in p.text)
        print(f"{Path(DOCX).parent.name}: 图段 {n_after} | 参考文献 {refs_after} 条 | 摘要 {len(NEW_ABSTRACT)} 字")


if __name__ == "__main__":
    main()
