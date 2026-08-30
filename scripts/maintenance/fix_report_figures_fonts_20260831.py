"""修复报告副本：删重复图、换新图、更图注、统一正文字体、表格居中。

用法: python scripts/maintenance/fix_report_figures_fonts_20260831.py
输入: D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx (原地更新)
依赖: python-docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DOCX_PATH = Path(r"D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx")
FIG_DIR = Path(r"D:/Project/厚粲杯/08_算法/output/06_正式实验/毫米波验证图_0831")

NEW_NOTES = {
    "图 11 毫米波融合心率与 ECG 金标准的对照":
        "注：图为 3 名佩戴 ECG 金标准参与者在 60 个探针前 30 s 窗口上，毫米波融合心率与 ECG 心率的 Bland-Altman 对照，偏倚 = −6.4 bpm，95% 一致性界限为 −28.6 至 15.8 bpm。",
    "图 12 心率融合估计机制":
        "注：A 为频域谱峰单独估计与完整链（融合后）心率相对 ECG 的散点；B 为两路估计绝对误差的分布。",
}


def _set_run_fonts(run, *, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def main() -> None:
    doc = Document(str(DOCX_PATH))
    paras = doc.paragraphs

    # 1) 删除 5.3.1 里重复插入的图 7（文字段前的孤立图段）
    removed = 0
    for i, p in enumerate(paras):
        blips = p._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
        if not blips or p.text.strip():
            continue
        if p.alignment is None and i + 1 < len(paras) and paras[i + 1].text.strip().startswith("61 名参与者的 109 个实验场次进入瞳孔分析"):
            p._element.getparent().remove(p._element)
            removed += 1
            print(f"已删除重复图段 [{i}]")
    print(f"删除重复图: {removed} 段")

    # 2) 替换图 11/12 图片并更新图注
    replaced = 0
    notes_fixed = 0
    for i, p in enumerate(paras):
        blips = p._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
        if not blips:
            continue
        nxt = paras[i + 1].text.strip() if i + 1 < len(paras) else ""
        if nxt.startswith("图 11") or nxt.startswith("图 12"):
            for run in list(p.runs):
                run._element.getparent().remove(run._element)
            fname = "MMWAVE_BLAND_ALTMAN_HR_2026-08-31.png" if nxt.startswith("图 11") else "MMWAVE_FUSION_MECHANISM_2026-08-31.png"
            p.add_run().add_picture(str(FIG_DIR / fname), width=Cm(13.5))
            p.alignment = 1  # CENTER
            replaced += 1
        if nxt in NEW_NOTES:
            note_para = paras[i + 2]
            for run in list(note_para.runs):
                run.text = ""
            note_para.runs[0].text = NEW_NOTES[nxt] if note_para.runs else None
            if not note_para.runs:
                _set_run_fonts(note_para.add_run(NEW_NOTES[nxt]), size_pt=9)
            else:
                _set_run_fonts(note_para.runs[0], size_pt=9)
            notes_fixed += 1
    print(f"替换图片: {replaced} 张；更新图注: {notes_fixed} 条")

    # 3) 统一 Normal 样式字体：中文宋体、西文 Times New Roman
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")
    print("Normal 样式字体已统一为 宋体 + Times New Roman")

    # 4) 表格居中
    centered = 0
    for tbl in doc.tables:
        tblPr = tbl._tbl.tblPr
        jc = tblPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            tblPr.append(jc)
        jc.set(qn("w:val"), "center")
        centered += 1
    print(f"表格居中: {centered} 张")

    doc.save(str(DOCX_PATH))
    print(f"已保存: {DOCX_PATH}")


if __name__ == "__main__":
    main()
