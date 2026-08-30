"""在 WPS 报告副本上填充 5.3 近红外瞳孔结果并修正 5.2 遗留术语。

用法: python scripts/maintenance/fill_report_53_20260831.py
输入: D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx (原地更新)
图源: D:/Project/厚粲杯/11_数据/_FormalAnalysis/NIR/12_pipeline_validation/figures/publication/
依赖: python-docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DOCX_PATH = Path(r"D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx")
FIG_DIR = Path(r"D:/Project/厚粲杯/11_数据/_FormalAnalysis/NIR/12_pipeline_validation/figures/publication")

FIGURES = {
    "5.3.1 瞳孔数据质量与可用范围": (
        "Figure09_data_quality_coverage.png",
        "图 7 瞳孔数据质量与覆盖",
        "注：呈现 109 个实验场次的瞳孔有效帧比例、窗口内部时间覆盖与双眼/单眼轨道可用性分布。",
    ),
    "5.3.2 视觉条件、任务进程与瞳孔变化": (
        "Figure06_visual_PLR_controls.png",
        "图 8 视觉刺激属性与瞳孔变化",
        "注：A 为不同刺激大小与任务类型下试次前中心化瞳孔的中位数；B、C、D 分别为瞳孔与当前刺激数字相对亮度、RMS 对比度及刺激可见面积的分箱散点。",
    ),
    "5.3.3 瞳孔与行为及思维探针的关系": (
        "Figure05_probe_states_trajectories.png",
        "图 9 探针状态与瞳孔轨迹",
        "注：呈现不同 Q1 注意内容与 Q2 警觉程度下探针前瞳孔几何平均直径的时间轨迹。",
    ),
}

# None 表示图占位（图、图题、图注三段）；5.3.3 末段图 10 单独插入
BODY = {
    "5.3.1 瞳孔数据质量与可用范围": [
        "61 名参与者的 109 个实验场次进入瞳孔分析。试次层面，80% 的试次瞳孔有效帧比例不低于 0.9，场次有效帧比例中位数为 1.000；窗口内部时间覆盖中位数为 0.960。左右眼分别保留测量，约六成试次以双眼同时可观测为主。八个候选直接几何指标中，以瞳孔几何平均直径与硬瞳孔面积比例作为主要指标，等效直径、椭圆长短轴、轮廓面积、椭圆面积与软瞳孔面积比例六个指标作为敏感性指标保留于附录。",
        None,
    ],
    "5.3.2 视觉条件、任务进程与瞳孔变化": [
        "为控制屏幕刺激引起的瞳孔光反射，按 9 类水果与 3 种尺寸组合的 27 种视觉条件计算数字相对亮度、RMS 对比度与刺激可见面积；试次模型同时纳入当前与前一试次的视觉属性以及任务时间、瞳孔有效帧比例、窗口覆盖与双眼帧比例等协变量。瞳孔对刺激尺寸、亮度与对比度的变化关系见图 8。",
        None,
        "试次级分析采用线性混合效应模型（正确 Go 反应时）与按参与者聚类的二项广义估计方程（Go 遗漏、No-Go 误按），瞳孔效应分解为参与者间平均水平与参与者内相对自身稳定水平的波动两部分，后者经参与者内标准化。正确 Go 反应时模型含 69,124 个试次，Go 遗漏模型含 70,649 个试次，No-Go 误按模型含 8,802 个试次，均来自 61 名参与者的 109 个实验场次。",
        "结果显示，参与者平均瞳孔几何直径每高 1 个标准差，正确 Go 反应时短 30.63 ms（95% CI [−51.82, −9.43]），调整视觉与质量协变量后为 −29.90 ms（95% CI [−51.27, −8.52]）；参与者内瞳孔相对波动的对应系数的 95% CI 包含 0。Go 遗漏方面，参与者内瞳孔相对波动每高 1 个标准差对应更低的遗漏对数发生比，B = −0.053（95% CI [−0.091, −0.015]），参与者间平均瞳孔的对应系数为 B = −0.628（95% CI [−0.880, −0.375]），调整后分别为 −0.050（95% CI [−0.090, −0.010]）与 −0.651（95% CI [−0.914, −0.388]）。No-Go 误按的参与者内与参与者间系数的 95% CI 均包含 0。总体上，更大的瞳孔对应更快的正确反应与更少的 Go 遗漏，且该关系在考虑视觉刺激属性与数据质量后保持稳定。完整参数与调整前后对照见附录 B。",
    ],
    "5.3.3 瞳孔与行为及思维探针的关系": [
        "探针层分析采用探针前 30 s 瞳孔指标，共 1,968 个探针观测，来自 61 名参与者的 109 个实验场次。以完全专注为参照类别的多项 Logistic 回归显示，瞳孔几何平均直径的参与者间项对“大脑空白”相对“完全专注”的系数为 B = 0.450（95% CI [0.097, 0.803]），硬瞳孔面积比例的对应系数为 B = 0.552（95% CI [0.184, 0.920]）；“在任务上没想目标”与“走神”相对参照类别的系数以及全部参与者内项的 95% CI 均包含 0。Q2 警觉程度的有序累积 Logit 模型中，两个主要指标全部系数的 95% CI 均包含 0。",
        None,
        "Block 层分析以任务区块为观测单位，共 214 个区块观测，来自 60 名参与者的 108 个实验场次。线性混合效应模型显示，瞳孔几何平均直径的参与者内相对项与 No-Go 误按率正相关，B = 0.089（95% CI [0.013, 0.164]），硬瞳孔面积比例对应 B = 0.080（95% CI [0.008, 0.153]）；两指标的参与者间项 95% CI 均包含 0。硬瞳孔面积比例的参与者间项与 Block 内反应时 Theil-Sen 斜率正相关，B = 0.022（95% CI [0.005, 0.039]）。辨别力 d′的对数线性估计、反应标准 c、β、正确 Go 反应时的均值、中位数、标准差、变异系数与 Go 遗漏率的各项系数 95% CI 均包含 0。总体上，区块内瞳孔相对增大与更高的 No-Go 误按率相关，参与者间平均瞳孔更大与“大脑空白”的探针报告相关。敏感性指标与全部 Block 层效应的结果见附录 B。",
        None,
    ],
}

FIG10 = (
    "Figure10_robustness_models.png",
    "图 10 瞳孔效应估计的稳健性",
    "注：呈现主要瞳孔指标在未调整与视觉调整模型中的效应估计与 95% CI。",
)


def _set_run_fonts(run, *, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_text_para(anchor, text: str) -> None:
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Pt(21)
    pf.line_spacing = 1.5
    _set_run_fonts(p.add_run(text))


def _add_figure_para(anchor, fname: str, caption: str, note: str) -> None:
    p_img = anchor.insert_paragraph_before()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(FIG_DIR / fname), width=Cm(13.5))
    p_cap = anchor.insert_paragraph_before()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_fonts(p_cap.add_run(caption), bold=True)
    p_note = anchor.insert_paragraph_before()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_fonts(p_note.add_run(note), size_pt=9)


def main() -> None:
    doc = Document(str(DOCX_PATH))

    # 1) 修正 5.2 旧术语与表编号引用
    fixed = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if "无时序歧义" in run.text or "见表 5-1" in run.text:
                run.text = (
                    run.text.replace("无时序歧义遗漏和时序歧义遗漏", "真遗漏和预判遗漏")
                    .replace("无时序歧义遗漏率均未出现", "真遗漏率均未出现")
                    .replace("时序歧义遗漏率在 B2", "预判遗漏率在 B2")
                    .replace("无时序歧义 Go 遗漏率", "真遗漏率")
                    .replace("见表 5-1", "见表 3")
                )
                fixed += 1
    print(f"术语/编号修正段落数: {fixed}")

    # 2) 删除 5.3.1 下的旧测试插入段
    for p in list(doc.paragraphs):
        if "【测试插入】" in p.text:
            p._element.getparent().remove(p._element)
            print("已删除测试插入段")

    # 3) 定位 5.3 小节标题与各节结束锚点
    headings: dict[str, object] = {}
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in BODY or t in {"5.4 毫米波胸部微动结果"}:
            headings[t] = p
    required = set(BODY) | {"5.4 毫米波胸部微动结果"}
    if not required.issubset(headings):
        raise RuntimeError(f"未找到全部锚点标题，缺失: {required - set(headings)}")

    anchors = {
        "5.3.1 瞳孔数据质量与可用范围": headings["5.3.2 视觉条件、任务进程与瞳孔变化"],
        "5.3.2 视觉条件、任务进程与瞳孔变化": headings["5.3.3 瞳孔与行为及思维探针的关系"],
        "5.3.3 瞳孔与行为及思维探针的关系": headings["5.4 毫米波胸部微动结果"],
    }

    # 4) 逐节插入内容（锚点前正序插入，保证顺序）
    for heading, anchor in anchors.items():
        for item in BODY[heading]:
            if item is None:
                _add_figure_para(anchor, *FIGURES[heading])
            else:
                _add_text_para(anchor, item)

    # 5) 5.3.3 末尾插入图 10
    _add_figure_para(headings["5.4 毫米波胸部微动结果"], *FIG10)

    doc.save(str(DOCX_PATH))
    print(f"已保存: {DOCX_PATH}")


if __name__ == "__main__":
    main()
