"""在 WPS 报告副本上填充 5.6 多模态增量与互补信息（含表 7 三线表）。

用法: python scripts/maintenance/fill_report_56_20260831.py
输入: D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx (原地更新)
依赖: python-docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DOCX_PATH = Path(r"D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx")

BODY = {
    "5.6.1 共同样本与基准模型": [
        "多模态比较在四模态共同可用的探针观测上进行，共 61 名参与者、109 个实验场次、2,180 个探针观测。行为、NIR 与 RGB 特征覆盖全部 2,180 个观测；毫米波有 6 个场次（120 个观测）的源数据结构性缺失，特征按缺失处理。Q1 注意内容的类别分布为 1,493、430、93 与 164，Q2 警觉程度为 39、318、1,004 与 819。所有组合以同一探针前 30 s 窗口、同一留一参与者交叉验证划分评估，标准化、缺失值插补与特征筛选只在训练折内拟合。",
        "以近期行为特征构成的基准模型（M0）在留一参与者交叉验证下，Q1 的 ROC-AUC 为 0.668、平衡准确率为 0.475、宏平均 F1 为 0.389；Q2 的 ROC-AUC 为 0.593、平衡准确率为 0.367、宏平均 F1 为 0.230。随机森林基准的对应 ROC-AUC 分别为 0.620 与 0.566。",
    ],
    "5.6.2 单模态、双模态与完整多模态增量": [
        "在基准之上依次加入 NIR、毫米波与 RGB 特征，比较三个单模态、三个双模态与完整三模态组合共八种模型（表 7）。相对基准的逐折配对差异以折为观测进行参与者级重抽样获得 95% CI。",
        "TABLE7",
        "Q1 上，各组合相对 M0 的 ROC-AUC 差异为 −0.011 至 +0.012，95% CI 均包含 0；平衡准确率上，加入 RGB 的组合呈小幅正增量（M3 +0.013，95% CI [0.005, 0.023]；M5 +0.017，95% CI [0.006, 0.028]），其余组合的 95% CI 包含 0；对数损失上，加入毫米波的组合呈小幅变差（M2 +0.015，95% CI [0.003, 0.028]；M7 +0.038，95% CI [0.002, 0.087]）。Q2 上，各组合相对 M0 的 ROC-AUC 差异均为负向且 95% CI 包含 0；宏平均 F1 上，含毫米波的组合呈正增量（M2 +0.038，95% CI [0.007, 0.066]；M4 +0.034，95% CI [0.004, 0.062]；M5 +0.035，95% CI [0.007, 0.065]）。",
        "随机森林补充模型的结果与逻辑回归不完全一致：Q1 上加入传感模态的组合一致改善对数损失（M7 相对 M0 为 −0.076，95% CI [−0.106, −0.047]），NIR+毫米波、毫米波+RGB 与完整组合的 ROC-AUC 正增量 95% CI 不含 0（+0.039 至 +0.043）；Q2 上未出现稳定增量。总体上，传感模态在行为基准之上的预测增量总体较小，RGB 对 Q1 平衡准确率的小幅正增量与毫米波对 Q2 宏 F1 的正增量是仅有的 95% CI 不含 0 的主模型增量，其余组合与指标的增量 95% CI 均包含 0。",
    ],
    "5.6.3 模态互补、平均边际贡献与补充 AI 融合": [
        "以全部 6 种加入顺序平均的边际贡献概括各模态的预测信息贡献。Q1 的 ROC-AUC 上，NIR、毫米波与 RGB 的平均边际贡献分别为 +0.005、−0.010 与 +0.010，平衡准确率上 RGB 为 +0.012；Q2 的 ROC-AUC 上分别为 −0.003、−0.013 与 +0.003。三种传感模态在已有信息基础上未呈现大幅联合增益，模态间信息以小幅叠加为主，未发现稳定的互补放大。该贡献仅描述预测信息，不作因果解释。",
        "非线性补充模型的相应结果见 5.6.2。",
    ],
}

TABLE7_HEADER = ["组合", "Q1 AUC", "Q1 平衡准确率", "Q1 宏 F1", "Q2 AUC", "Q2 平衡准确率", "Q2 宏 F1"]
TABLE7_ROWS = [
    ["M0 行为基准", "0.668", "0.475", "0.389", "0.593", "0.367", "0.230"],
    ["M1 +NIR", "0.677", "0.475", "0.397", "0.592", "0.370", "0.235"],
    ["M2 +毫米波", "0.657", "0.472", "0.383", "0.576", "0.377", "0.268"],
    ["M3 +RGB", "0.676", "0.488", "0.389", "0.591", "0.393", "0.252"],
    ["M4 NIR+毫米波", "0.660", "0.476", "0.376", "0.572", "0.365", "0.264"],
    ["M5 NIR+RGB", "0.680", "0.491", "0.406", "0.588", "0.395", "0.265"],
    ["M6 毫米波+RGB", "0.670", "0.483", "0.387", "0.584", "0.380", "0.265"],
    ["M7 完整三模态", "0.674", "0.485", "0.391", "0.580", "0.376", "0.265"],
]
TABLE7_NOTE = "注：表中为规则化多项 Logistic 回归结果，逐折等权平均。AUC 为 one-vs-rest 宏平均，Q1 有 11 个测试折、Q2 有 7 个测试折因该参与者类别不全而无法定义，表中为该指标可定义折的均值。"


def _set_run_fonts(run, *, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_text_para(anchor, text: str) -> None:
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Pt(21)
    pf.line_spacing = 1.5
    _set_run_fonts(p.add_run(text))


def _header_rule(table) -> None:
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tcPr.append(borders)
        b = borders.find(qn("w:bottom"))
        if b is None:
            b = OxmlElement("w:bottom")
            borders.append(b)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")


def _add_table7(anchor) -> None:
    doc = anchor._parent
    tbl = doc.add_table(rows=len(TABLE7_ROWS) + 1, cols=7, width=Cm(14.0))
    try:
        tbl.style = anchor.part.document.styles["三线表"]
    except KeyError:
        pass
    for j, h in enumerate(TABLE7_HEADER):
        tbl.rows[0].cells[j].text = h
    for i, row in enumerate(TABLE7_ROWS, start=1):
        for j, val in enumerate(row):
            tbl.rows[i].cells[j].text = val
    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    _set_run_fonts(run, size_pt=9, bold=(ri == 0))
    _header_rule(tbl)
    anchor._element.addprevious(tbl._element)


def main() -> None:
    doc = Document(str(DOCX_PATH))

    headings: dict[str, object] = {}
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in BODY or t == "6 分析与讨论":
            headings[t] = p
    required = set(BODY) | {"6 分析与讨论"}
    if not required.issubset(headings):
        raise RuntimeError(f"缺失标题: {required - set(headings)}")

    anchors = {
        "5.6.1 共同样本与基准模型": headings["5.6.2 单模态、双模态与完整多模态增量"],
        "5.6.2 单模态、双模态与完整多模态增量": headings["5.6.3 模态互补、平均边际贡献与补充 AI 融合"],
        "5.6.3 模态互补、平均边际贡献与补充 AI 融合": headings["6 分析与讨论"],
    }

    for heading, anchor in anchors.items():
        for item in BODY[heading]:
            if item == "TABLE7":
                _add_table7(anchor)
                p_cap = anchor.insert_paragraph_before()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_run_fonts(p_cap.add_run("表 7 八种模态组合的预测性能（留一参与者交叉验证，逐折等权）"), bold=True)
                p_note = anchor.insert_paragraph_before()
                p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_run_fonts(p_note.add_run(TABLE7_NOTE), size_pt=9)
            else:
                _add_text_para(anchor, item)

    doc.save(str(DOCX_PATH))
    print(f"已保存: {DOCX_PATH}")


if __name__ == "__main__":
    main()
