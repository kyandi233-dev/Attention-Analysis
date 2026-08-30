"""在 WPS 报告副本上填充 5.5 RGB 可见行为结果与图 13-16。

用法: python scripts/maintenance/fill_report_55_20260831.py
输入: D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx (原地更新)
图源: D:/Project/厚粲杯/11_数据/_FormalAnalysis/RGB/21_analysis_tables_5.5/figures/
依赖: python-docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DOCX_PATH = Path(r"D:/Project/厚粲杯/08_算法/docs/交付/0827报告v1_填充版_20260831.docx")
FIG_DIR = Path(r"D:/Project/厚粲杯/11_数据/_FormalAnalysis/RGB/21_analysis_tables_5.5/figures")

FIGURES = {
    "5.5.1 整体运动与画面亮度变化": (
        "fig_55_1_motion_exposure.png",
        "图 13 RGB 运动能量与曝光变化的任务进程",
        "注：A、B 为运动能量与曝光变化的 B1/B2 配对；C、D 为 Block×cycle 轨迹（参与者均值 ± 标准误）。",
    ),
    "5.5.2 姿态与方向性运动": (
        "fig_55_2_pose_direction.png",
        "图 14 RGB 姿态方向候选的分布",
        "注：左右、上下方向运动与复合前后方向候选（图像内无量纲）在探针前窗口的分布。",
    ),
    "5.5.3 眨眼候选事件": (
        "fig_55_3_blink_events.png",
        "图 15 RGB 眨眼候选事件的分布",
        "注：A、B、C 分别为眨眼候选事件场次级频率、时长与间隔分布；数据覆盖见图 16。",
    ),
}

BODY = {
    "5.5.1 整体运动与画面亮度变化": [
        "RGB 可见行为分析以 116 个实验场次的正式队列为口径，其中 115 个场次具有可用视频记录（1 个场次无 RGB 源），共形成 2,320 个探针前 30 s 窗口观测。身体运动、曝光变化、姿态方向与眨眼候选按探针前窗口和 Block 内时间进程两个尺度汇总，探针前窗口严格排除探针锚定试次及其后的帧。",
        "身体运动可观测帧比例的场次均值为 0.925（SD = 0.003），曝光变化与运动能量保持分轨计算。Block×cycle 的参与者聚类 GEE 显示，运动能量与曝光变化的 Block、cycle 及交互项系数的 95% CI 均包含 0，两个任务区块内的整体运动与画面亮度变化保持稳定（图 13）。运动能量、曝光变化与 Q1 注意内容、Q2 警觉程度及近期行为指标的关联中，各系数 95% CI 均包含 0；与毫米波运动代理、心率与呼吸频率的窗口级对照同样未出现 95% CI 不含 0 的关联。",
        None,
    ],
    "5.5.2 姿态与方向性运动": [
        "肩部关键点可观测帧比例的场次均值为 0.993（最小 0.819）。左右与上下方向的可见运动以图像内方向变化描述；前后方向由深度方向、肩宽与人体面积变化综合为无量纲的复合方向候选，不解释为物理位移。探针窗口的方向特征与 Q1 的多项 Logistic 回归显示，左右方向运动的窗口中位数与“在任务上没想目标”相对“完全专注”的系数为 B = −0.119，95% CI [−0.219, −0.019]；上下方向运动与“大脑空白”相对“完全专注”的系数为 B = −0.220，95% CI [−0.422, −0.018]；其余类别的系数 95% CI 均包含 0。方向特征与 Q2 的系数 95% CI 均包含 0。前后方向候选与正确 Go RT 中位数的系数为 B = −0.010，95% CI [−0.019, −0.001]，量级很小；方向特征与其余行为指标及毫米波指标的对照未出现 95% CI 不含 0 的关联。",
        None,
    ],
    "5.5.3 眨眼候选事件": [
        "左右眼可观测帧比例的场次均值均为 0.999（114 个场次），双眼一致性帧比例的场次均值为 0.960（SD = 0.059，最小 0.443）。110 个场次形成眨眼候选事件表，每场次候选事件数中位数为 270（范围 1–1,200），候选频率场次均值为 13.0 次/分（SD = 11.0），候选时长中位数为 65 ms（单帧事件的 65 ms 为名义值），事件间隔中位数为 1,015 ms。独立事件标注效度尚未建立，以下均按算法定义的候选事件表述。",
        "Block×cycle 的 GEE 显示，眨眼候选事件率在 B2 中更高（B = 2.20，95% CI [0.47, 3.94]），并随 Block 内 cycle 推进而升高（B = 0.51，95% CI [0.19, 0.83]），Block×cycle 交互项的 95% CI 包含 0。眨眼候选事件率与 Q2 警觉程度呈负向关系，B = −0.244，95% CI [−0.416, −0.072]，即候选率较高的探针窗口对应较低的警觉等级。眨眼候选事件率与窗口行为指标的关系显示，候选率与 Go 遗漏率（B = 1.56，95% CI [0.35, 2.78]）和 No-Go 误按率（B = 0.60，95% CI [0.09, 1.11]）正相关，与 d′负相关（B = −1.10，95% CI [−1.92, −0.27]）。眨眼候选事件率与毫米波运动代理正相关（B = 0.625，95% CI [0.145, 1.105]），与毫米波心率、呼吸频率的对照未出现 95% CI 不含 0 的关联。",
        None,
    ],
}

FIG16 = (
    "fig_55_4_coverage.png",
    "图 16 RGB 各分析轨道的场次级可观测覆盖",
    "注：运动、曝光、肩部、左右眼与双眼一致性六轨的场次级可观测帧比例。",
)


def _set_run_fonts(run, *, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_text_para(anchor, text: str) -> None:
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Pt(21)
    pf.line_spacing = 1.5
    _set_run_fonts(p.add_run(text))


def _add_figure_para(anchor, fname: str, caption: str, note: str) -> None:
    p_img = anchor.insert_paragraph_before()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(FIG_DIR / fname), width=Cm(13.5))
    p_cap = anchor.insert_paragraph_before()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_fonts(p_cap.add_run(caption), bold=True)
    p_note = anchor.insert_paragraph_before()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_fonts(p_note.add_run(note), size_pt=9)


def main() -> None:
    doc = Document(str(DOCX_PATH))

    headings: dict[str, object] = {}
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in BODY or t == "5.6 多模态增量与互补信息":
            headings[t] = p
    required = set(BODY) | {"5.6 多模态增量与互补信息"}
    if not required.issubset(headings):
        raise RuntimeError(f"缺失标题: {required - set(headings)}")

    anchors = {
        "5.5.1 整体运动与画面亮度变化": headings["5.5.2 姿态与方向性运动"],
        "5.5.2 姿态与方向性运动": headings["5.5.3 眨眼候选事件"],
        "5.5.3 眨眼候选事件": headings["5.6 多模态增量与互补信息"],
    }

    for heading, anchor in anchors.items():
        for item in BODY[heading]:
            if item is None:
                _add_figure_para(anchor, *FIGURES[heading])
            else:
                _add_text_para(anchor, item)

    _add_figure_para(headings["5.6 多模态增量与互补信息"], *FIG16)

    doc.save(str(DOCX_PATH))
    print(f"已保存: {DOCX_PATH}")


if __name__ == "__main__":
    main()
